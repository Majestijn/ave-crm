from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_full_script():
    doc = Document()

    # --- Styles ---
    # Title
    title = doc.add_heading('Volledig Uitgeschreven Script: Eindpresentatie AVE CRM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Spreektaal - Klaar om voor te lezen of te oefenen')
    run.italic = True
    
    doc.add_paragraph() # Spacer

    # Helper for adding script sections
    def add_slide_script(slide_num, title, text, cues=None):
        # Header for the slide
        h = doc.add_heading(f"Slide {slide_num}: {title}", level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

        # Optional cues (actions)
        if cues:
            p_cue = doc.add_paragraph()
            run_cue = p_cue.add_run(f"[ACTIE: {cues}]")
            run_cue.bold = True
            run_cue.italic = True
            run_cue.font.color.rgb = RGBColor(200, 0, 0) # Red for instructions

        # The spoken text
        p = doc.add_paragraph(text)
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(11)
        
        doc.add_paragraph() # Spacer

    # --- SCRIPT CONTENT ---

    add_slide_script(
        1, "Titel Slide",
        "Goedemorgen allemaal. Welkom bij mijn eindpresentatie.\n\n"
        "Mijn naam is Stijn van der Neut en de afgelopen 20 weken heb ik mij beziggehouden met de digitale transformatie van AVE Consultancy. "
        "Vandaag neem ik jullie mee in de reis van een klassieke, analoge werkwijze naar een modern, digitaal SaaS-platform. "
        "Ik vertel jullie niet alleen WAT ik heb gebouwd, maar vooral WAAROM, en hoe ik mijzelf tijdens dit proces heb ontwikkeld van student naar professional."
    )

    add_slide_script(
        2, "Agenda",
        "Om structuur te geven aan het verhaal, beginnen we bij de basis: de situatie zoals ik die aantrof. "
        "Daarna kijken we naar het probleem dat daaruit voortkwam en het onderzoek dat ik heb gedaan. "
        "Vervolgens duiken we de diepte in: ik laat jullie de oplossing zien en we bespreken twee technische hoogtepunten: Multi-Tenancy en AI. "
        "Ik sluit af met een persoonlijke reflectie op mijn leerproces en een blik op de toekomst.\n\n"
        "Vragen mogen tussendoor als ze dringend zijn, maar voor de flow bewaar ik ze het liefst voor het einde."
    )

    add_slide_script(
        3, "Situatieschets & Aanleiding",
        "Laten we teruggaan naar september. AVE Consultancy is een succesvol headhuntingbureau met de ambitie om te groeien. "
        "Maar als we onder de motorkap keken, zagen we dat de bedrijfsprocessen die ambitie niet konden bijbenen.\n\n"
        "De situatie was als volgt: informatie stond versnipperd. CV's stonden in mappen op Dropbox, klantgegevens in verschillende Excel-lijsten en communicatie zat vast in de mailboxen van individuele medewerkers. "
        "Er was geen centraal brein. Als Adriaan wilde weten: 'Welke kandidaten hebben we voorgesteld aan Klant X?', dan was dat een zoektocht van soms wel een uur."
    )

    add_slide_script(
        4, "Opdracht, Scope & Tijdsframe",
        "Dat moest anders. De opdracht die ik kreeg was helder, maar uitdagend: 'Ontwikkel een fundering voor de toekomst'. "
        "Niet zomaar een database, maar een SaaS-platform (Software as a Service) waarmee AVE niet alleen zelf kan werken, maar dat in de toekomst ook aan andere bureaus verkocht kan worden.\n\n"
        "Ik had 20 weken de tijd. We hebben daarom een strakke scope bepaald voor een MVP (Minimum Viable Product). "
        "De focus lag op de kern van het vak: Relaties beheren. Kandidaten, Klanten en de Opdrachten daartussen. "
        "Zaken als facturatie of een mobiele app hebben we bewust buiten beschouwing gelaten om kwaliteit te kunnen garanderen."
    )

    add_slide_script(
        5, "Probleemstelling",
        "Waarom was die oude situatie nu zo problematisch? \n"
        "Ten eerste: Inefficiëntie. Het handmatig verwerken van honderden CV's kostte letterlijk dagen werk.\n"
        "Ten tweede: Risico. We werken met persoonsgegevens. Excel-lijstjes heen en weer mailen is in 2026 echt niet meer AVG-proof.\n"
        "En ten derde: Gebrek aan inzicht. Zonder relaties in je data kun je niet sturen op cijfers. Je vaart blind."
    )

    add_slide_script(
        6, "Onderzoek (Build vs Buy)",
        "Als HBO-professional ga je niet meteen bouwen. Je gaat eerst analyseren. Moeten we dit wel zelf maken?\n\n"
        "Ik heb volgens de DSR-methode (Design Science Research) gekeken naar de markt. We zagen twee smaken:\n"
        "Aan de ene kant de Enterprise giganten zoals Bullhorn of Salesforce. Geweldig, maar extreem duur en complex om in te richten voor een klein bureau.\n"
        "Aan de andere kant systemen zoals Recruitee. Betaalbaar, maar die zijn gemaakt voor HR-afdelingen, niet voor bureaus die 'makelen' tussen partijen.\n\n"
        "De conclusie was duidelijk: Er is een 'gap' in de markt. Maatwerk was de enige manier om de specifieke werkwijze van AVE te ondersteunen én eigenaar te blijven van de data."
    )

    add_slide_script(
        7, "De Oplossing (Tech Stack)",
        "Dus zijn we gaan bouwen. Ik heb gekozen voor een robuuste, moderne tech stack.\n\n"
        "Aan de achterkant (Backend) draait Laravel (PHP). Dit is de wereldwijde standaard voor SaaS-applicaties: veilig en stabiel.\n"
        "Aan de voorkant (Frontend) zien de gebruikers een React applicatie. Dit zorgt voor die snelle, 'snappy' ervaring die je verwacht van moderne software, zonder dat de pagina steeds moet herladen.\n"
        "Voor de opslag van die duizenden CV's gebruiken we Cloudflare R2. Dat is net zo goed als Amazon S3, maar een stuk goedkoper en sneller."
    )

    add_slide_script(
        8, "Diepgang 1: Multi-Tenancy",
        "Dan nu de technische diepgang. Want hoe zorg je er in een SaaS-omgeving voor dat Klant A nooit de data van Klant B ziet?\n\n"
        "Ik heb gekozen voor een 'Database-per-Tenant' strategie. Dit is de meest veilige optie. \n"
        "Iedere klant die inlogt, krijgt zijn eigen, fysiek gescheiden database. \n"
        "Het systeem kijkt naar het domein, bijvoorbeeld 'klant-a.avecrm.nl', en weet dan: ik mag alléén verbinden met Database A.\n"
        "Zelfs als ik als programmeur een fout maak in de code, is het technisch onmogelijk om data van de verkeerde klant op te halen. Veiligheid 'by design' dus."
    )

    add_slide_script(
        9, "Diepgang 2: AI Bulk Import",
        "Het tweede technische hoogtepunt loste ons grootste probleem op: De historie. We hadden 3500 oude CV's in mapjes.\n\n"
        "Ik heb een AI-pipeline gebouwd met Google Gemini 3 Pro.\n"
        "Het werkt zo: Je sleept 100 CV's in het systeem. De server pakt ze op, en de AI 'leest' ze als een mens.\n"
        "Hij haalt de naam, e-mail, skills en werkervaring eruit en stopt dit netjes in de database.\n"
        "Wat vroeger 15 minuten per CV kostte aan typewerk, gebeurt nu in enkele seconden. Dit is de ware kracht van digitalisering.",
        cues="Als je een video/demo hebt, start die hier."
    )

    add_slide_script(
        10, "Persoonlijke Ontwikkeling (Veerkracht)",
        "Tijdens dit project ging niet alles vlekkeloos. En daar wil ik eerlijk over zijn.\n\n"
        "Halverwege de stage, in Sprint 4, crashte mijn ontwikkelomgeving. Omdat ik geen goede backups had, was ik een week werk kwijt.\n"
        "Mijn eerste reactie was paniek. Ik trok me terug, het zogenoemde 'oestergedrag'. Ik dacht: ik los dit wel alleen op.\n"
        "Maar ik leerde dat dat niet werkt. Ik heb het opgebiecht aan mijn begeleider. In plaats van boosheid, kreeg ik hulp.\n"
        "Ik heb diezelfde dag nog een geautomatiseerd backup-script geschreven. \n"
        "De les die ik meeneem: Fouten maken mag, zolang je erover communiceert en het oplost."
    )

    add_slide_script(
        11, "Persoonlijke Ontwikkeling (Professionaliteit)",
        "Als ik kijk naar de Stijn van 20 weken geleden, zie ik een afwachtende student. Ik vroeg: 'Wat moet ik doen?'\n\n"
        "Nu sta ik hier als professional. Ik wacht niet meer af, ik stel voor. \n"
        "Ik heb zelf de wekelijkse meetings opgezet, ik beheer de planning en ik adviseer Adriaan over technische keuzes.\n"
        "Zoals in de feedback van Hugo stond: Ik heb de rol gepakt van 'Strategisch Partner'."
    )

    add_slide_script(
        12, "Toekomstvisie",
        "En nu? De stage stopt, maar het product leeft.\n\n"
        "De MVP gaat live. We gaan het systeem nu intern gebruiken ('Dogfooding') om de laatste puntjes op de i te zetten.\n"
        "De volgende stap is de koppeling met Outlook, zodat ook de agenda's gesynchroniseerd zijn.\n"
        "En op de lange termijn staat de weg open om dit platform in de markt te zetten voor andere bureaus."
    )

    add_slide_script(
        13, "Conclusie",
        "Samenvattend: We zijn in 20 weken van een analoge chaos naar een gestructureerd, digitaal fundament gegaan.\n"
        "Er staat een veilig systeem, er is een slimme AI-oplossing en ik heb mijzelf ontwikkeld tot een zelfstandige developer.\n\n"
        "Ik wil Adriaan en mijn begeleiders bedanken voor het vertrouwen.\n"
        "Dit was mijn presentatie. Zijn er nog vragen?"
    )

    # Save
    filename = "Volledig_Script_Eindpresentatie_AVE_CRM.docx"
    doc.save(filename)
    print(f"Successfully generated '{filename}'")

if __name__ == "__main__":
    create_full_script()
