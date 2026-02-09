from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_click_script():
    doc = Document()

    # --- Styles ---
    # Title
    title = doc.add_heading('Klik-Script: Eindpresentatie AVE CRM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Met exacte [KLIK] momenten voor naadloze timing')
    run.italic = True
    
    doc.add_paragraph() # Spacer

    # Helper for adding script sections
    def add_section(slide_num, title, text_blocks):
        # Header
        h = doc.add_heading(f"Slide {slide_num}: {title}", level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        
        # Blocks
        for block in text_blocks:
            if block == "[KLIK]":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("--- [KLIK] NAAR VOLGENDE SLIDE ---")
                run.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0) # Red
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
            else:
                p = doc.add_paragraph(block)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)

        doc.add_paragraph() # Spacer between slides

    # --- SCRIPT CONTENT ---

    add_section(
        1, "Titel Slide",
        [
            "Goedemorgen allemaal. Welkom bij mijn eindpresentatie.",
            "Mijn naam is Stijn van der Neut en de afgelopen 20 weken heb ik mij beziggehouden met de digitale transformatie van AVE Consultancy.",
            "Vandaag neem ik jullie mee in de reis van een klassieke, analoge werkwijze naar een modern, digitaal SaaS-platform.",
            "Ik vertel jullie niet alleen WAT ik heb gebouwd, maar vooral WAAROM, en hoe ik mijzelf tijdens dit proces heb ontwikkeld van student naar professional.",
            "[KLIK]"
        ]
    )

    add_section(
        2, "Agenda",
        [
            "Om structuur te geven aan het verhaal, beginnen we bij de basis: de situatie zoals ik die aantrof.",
            "Daarna kijken we naar het probleem dat daaruit voortkwam en het onderzoek dat ik heb gedaan.",
            "Vervolgens duiken we de diepte in: ik laat jullie de oplossing zien en we bespreken twee technische hoogtepunten: Multi-Tenancy en AI.",
            "Ik sluit af met een persoonlijke reflectie op mijn leerproces en een blik op de toekomst.",
            "Vragen mogen tussendoor als ze dringend zijn, maar voor de flow bewaar ik ze het liefst voor het einde.",
            "[KLIK]"
        ]
    )

    add_section(
        3, "Situatieschets & Aanleiding",
        [
            "Laten we teruggaan naar september. AVE Consultancy is een succesvol headhuntingbureau met de ambitie om te groeien.",
            "Maar als we onder de motorkap keken, zagen we dat de bedrijfsprocessen die ambitie niet konden bijbenen.",
            "De situatie was als volgt: informatie stond versnipperd. CV's stonden in mappen op Dropbox, klantgegevens in verschillende Excel-lijsten en communicatie zat vast in de mailboxen van individuele medewerkers.",
            "Het gevolg was simpel maar pijnlijk: Er was geen centraal inzicht. Als Adriaan wilde weten: 'Welke kandidaten hebben we voorgesteld aan Klant X?', dan was dat een zoektocht van soms wel een uur.",
            "[KLIK]"
        ]
    )

    add_section(
        4, "Opdracht, Scope & Tijdsframe",
        [
            "Dat moest anders. De opdracht die ik kreeg was helder, maar uitdagend: 'Ontwikkel een fundering voor de toekomst'.",
            "Niet zomaar een database, maar een SaaS-platform (Software as a Service) waarmee AVE niet alleen zelf kan werken, maar dat in de toekomst ook aan andere bureaus verkocht kan worden.",
            "Ik had 20 weken de tijd. We hebben daarom een strakke scope bepaald voor een MVP (Minimum Viable Product).",
            "De focus lag op de kern van het vak: Relaties beheren. Kandidaten, Klanten en de Opdrachten daartussen.",
            "Zaken als facturatie of een mobiele app hebben we bewust buiten beschouwing gelaten om kwaliteit te kunnen garanderen.",
            "[KLIK]"
        ]
    )

    add_section(
        5, "Probleemstelling",
        [
            "Waarom was die oude situatie nu zo problematisch? Ik heb dit samengevat in drie punten.",
            "Ten eerste: Inefficiëntie. Het handmatig verwerken van honderden CV's kostte letterlijk dagen werk.",
            "Ten tweede: Risico. We werken met persoonsgegevens. Excel-lijstjes heen en weer mailen is in 2026 echt niet meer AVG-proof.",
            "En ten derde: Gebrek aan inzicht. Zonder relaties in je data kun je niet sturen op cijfers. Je vaart blind.",
            "[KLIK]"
        ]
    )

    add_section(
        6, "Onderzoek (Build vs Buy)",
        [
            "Als HBO-professional ga je niet meteen bouwen. Je gaat eerst analyseren. Moeten we dit wel zelf maken?",
            "Ik heb volgens de DSR-methode gekeken naar de markt en zag twee smaken:",
            "Optie A: De Enterprise giganten zoals Bullhorn of Salesforce. Geweldig, maar extreem duur en complex om in te richten voor een klein bureau.",
            "Optie B: Systemen zoals Recruitee. Betaalbaar, maar die zijn gemaakt voor HR-afdelingen, niet voor bureaus die 'makelen' tussen partijen.",
            "De conclusie was duidelijk: Er is een 'gap' in de markt. Maatwerk was de enige manier om de specifieke werkwijze van AVE te ondersteunen én eigenaar te blijven van de data.",
            "[KLIK]"
        ]
    )

    add_section(
        7, "De Oplossing (Tech Stack)",
        [
            "Dus zijn we gaan bouwen. Ik heb gekozen voor een robuuste, moderne tech stack.",
            "Aan de achterkant draait Laravel (PHP). Dit is de wereldwijde standaard voor SaaS-applicaties: veilig en stabiel.",
            "Aan de voorkant zien de gebruikers een React applicatie. Dit zorgt voor die snelle, 'snappy' ervaring die je verwacht van moderne software.",
            "Voor de opslag van die duizenden CV's gebruiken we Cloudflare R2. Dat is net zo goed als Amazon S3, maar een stuk goedkoper en sneller.",
            "[KLIK]"
        ]
    )

    add_section(
        8, "Diepgang 1: Multi-Tenancy",
        [
            "Dan nu de technische diepgang. Want hoe zorg je er in een SaaS-omgeving voor dat Klant A nooit de data van Klant B ziet?",
            "Ik heb gekozen voor een 'Database-per-Tenant' strategie. Dit is de meest veilige optie.",
            "Iedere klant die inlogt, krijgt zijn eigen, fysiek gescheiden database.",
            "Het systeem kijkt naar het domein, bijvoorbeeld 'klant-a.avecrm.nl', en weet dan: ik mag alléén verbinden met Database A.",
            "Zelfs als ik als programmeur een fout maak in de code, is het technisch onmogelijk om data van de verkeerde klant op te halen. Veiligheid 'by design' dus.",
            "[KLIK]"
        ]
    )

    add_section(
        9, "Diepgang 2: AI Bulk Import",
        [
            "Het tweede technische hoogtepunt loste ons grootste probleem op: De historie. We hadden 3500 oude CV's in mapjes.",
            "Ik heb een AI-pipeline gebouwd met Google Gemini 3 Pro.",
            "Het werkt zo: Je sleept 100 CV's in het systeem. De server pakt ze op, en de AI 'leest' ze als een mens.",
            "Hij haalt de naam, e-mail, skills en werkervaring eruit en stopt dit netjes in de database.",
            "Wat vroeger 15 minuten per CV kostte aan typewerk, gebeurt nu in enkele seconden. Dit is de ware kracht van digitalisering.",
            "[KLIK]"
        ]
    )

    add_section(
        10, "Reflectie: Veerkracht",
        [
            "Tijdens dit project ging niet alles vlekkeloos. En daar wil ik eerlijk over zijn.",
            "Halverwege de stage, in Sprint 4, crashte mijn ontwikkelomgeving. Omdat ik geen goede backups had, was ik een week werk kwijt.",
            "Mijn eerste reactie was paniek. Ik trok me terug, het zogenoemde 'oestergedrag'. Ik dacht: ik los dit wel alleen op.",
            "Maar ik leerde dat dat niet werkt. Ik heb het opgebiecht aan mijn begeleider. In plaats van boosheid, kreeg ik hulp.",
            "Ik heb diezelfde dag nog een geautomatiseerd backup-script geschreven. De les die ik meeneem: Fouten maken mag, zolang je erover communiceert en het oplost.",
            "[KLIK]"
        ]
    )

    add_section(
        11, "Reflectie: Professionaliteit",
        [
            "Als ik kijk naar de Stijn van 20 weken geleden, zie ik een afwachtende student. Ik vroeg: 'Wat moet ik doen?'.",
            "Nu sta ik hier als professional. Ik wacht niet meer af, ik stel voor.",
            "Ik heb zelf de wekelijkse meetings opgezet, ik beheer de planning en ik adviseer Adriaan over technische keuzes.",
            "Zoals in de feedback van Hugo stond: Ik heb de rol gepakt van 'Strategisch Partner'.",
            "[KLIK]"
        ]
    )

    add_section(
        12, "Toekomstvisie",
        [
            "En nu? De stage stopt, maar het product leeft.",
            "De MVP gaat live. We gaan het systeem nu intern gebruiken ('Dogfooding') om de laatste puntjes op de i te zetten.",
            "De volgende stap is de koppeling met Outlook, zodat ook de agenda's gesynchroniseerd zijn.",
            "En op de lange termijn staat de weg open om dit platform in de markt te zetten voor andere bureaus.",
            "[KLIK]"
        ]
    )

    add_section(
        13, "Conclusie",
        [
            "Samenvattend: We zijn in 20 weken van een analoge chaos naar een gestructureerd, digitaal fundament gegaan.",
            "Er staat een veilig systeem, er is een slimme AI-oplossing en ik heb mijzelf ontwikkeld tot een zelfstandige developer.",
            "Ik wil Adriaan en mijn begeleiders bedanken voor het vertrouwen.",
            "Dit was mijn presentatie. Zijn er nog vragen?"
        ]
    )

    # Save
    filename = "Volledig_Script_Met_Klikmomenten.docx"
    doc.save(filename)
    print(f"Successfully generated '{filename}'")

if __name__ == "__main__":
    create_click_script()
