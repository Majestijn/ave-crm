from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Stijlen definiëren
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Titel
    title = doc.add_heading('Business Case: AVE CRM Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header informatie
    info = doc.add_paragraph()
    info.add_run('Betreft:').bold = True
    info.add_run(' Software analyse t.b.v. financieel advies en waardebepaling\n')
    info.add_run('Datum:').bold = True
    info.add_run(' 30 januari 2026\n')
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_section(title_text, color):
        h = doc.add_heading(title_text, level=1)
        for run in h.runs:
            run.font.color.rgb = color
            run.font.size = Pt(14)

    # Kleur: Donkerrood uit het project (#800400)
    brand_color = RGBColor(128, 4, 0)

    # Sectie 1: Wat is AVE CRM?
    add_section('1. Wat is AVE CRM?', brand_color)
    p1 = doc.add_paragraph(
        "AVE CRM is een modern, cloud-based softwareplatform specifiek ontwikkeld voor de werving- en selectiebranche. "
        "Het systeem digitaliseert en automatiseert het volledige proces van kandidaat-bemiddeling: van het importeren "
        "van CV's tot het matchen van kandidaten op opdrachten bij klanten."
    )
    
    doc.add_paragraph(
        "In tegenstelling tot standaardpakketten is AVE CRM gebouwd met de nieuwste technologieën (AI, Cloud) en "
        "volledig afgestemd op de specifieke werkwijze van AVE, met een sterke focus op snelheid en privacy."
    )

    # Sectie 2: Waarde voor AVE Services
    add_section('2. Waarde voor AVE Services (Interne Business Case)', brand_color)
    
    bullets = [
        ("Efficiëntieslag door AI:", "Gebruik van Google Gemini & Vertex AI om automatisch CV's uit te lezen. Bespaart recruiters uren aan handmatig invoerwerk per week."),
        ("Professionalisering:", "Integratie met Microsoft 365 (Agenda & Mail) en een moderne interface voor een sterke marktuitstraling."),
        ("Centrale Data & Inzicht:", "Eliminatie van versnipperde Excel-lijstjes; alle stuurinformatie (KPI's) is direct en centraal inzichtelijk."),
        ("AVG/GDPR Compliance:", "Privacy by Design opzet met veilige dataverwerking binnen Europa, essentieel voor gevoelige kandidaatgegevens.")
    ]

    for bold_text, normal_text in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(f" {normal_text}")

    # Sectie 3: Commerciële Potentie
    add_section('3. Commerciële Potentie (Doorverkoop aan Derden)', brand_color)
    
    bullets_comm = [
        ("Multi-Tenant Architectuur:", "Technisch gebouwd om eenvoudig nieuwe, afgeschermde omgevingen voor externe klanten op te starten (SaaS-model)."),
        ("Generieke Marktbehoefte:", "De oplossing is niet AVE-specifiek maar lost universele problemen in de recruitmentsector op."),
        ("Toekomstbestendig:", "Gebouwd met de nieuwste standaarden (Laravel 12, React 19). Dit garandeert een lange technische levensduur en hoge marktwaarde."),
        ("Scalair Verdienmodel:", "Uitstekend geschikt voor een licentiemodel (MRR - Monthly Recurring Revenue), wat zorgt voor een voorspelbare waardestijging.")
    ]

    for bold_text, normal_text in bullets_comm:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_text)
        run.bold = True
        p.add_run(f" {normal_text}")

    # Voetnoot
    doc.add_paragraph()
    footer = doc.add_paragraph("Vertrouwelijk document - Enkel voor intern gebruik en financieel advies.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.italic = True

    doc.save('AVE_CRM_Business_Case.docx')
    print("Document succesvol gegenereerd: AVE_CRM_Business_Case.docx")

if __name__ == "__main__":
    create_document()
