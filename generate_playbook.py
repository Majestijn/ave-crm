from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_playbook():
    doc = Document()

    # --- Styles Setup ---
    # Title
    title = doc.add_heading('Draaiboek Eindpresentatie: AVE CRM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Datum: 19 Januari 2026 | Spreker: Stijn van der Neut | Duur: ~20 min + Vragen')
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph() # Spacer

    # --- Section 1: Voorbereiding ---
    doc.add_heading('1. Voorbereiding & Checklist (5 min voor start)', level=1)
    
    checklist = [
        "Laptop aansluiten op scherm (HDMI/USB-C).",
        "Presentatie openen in 'Presenter View' (zodat je notities ziet).",
        "Zorg dat de demo-omgeving (localhost) draait voor het geval er vragen zijn.",
        "Glas water klaarzetten.",
        "Telefoon op 'Niet storen'.",
        "Ademhaling check: Rustig in, rustig uit. Je bent de expert van dit project."
    ]
    
    for item in checklist:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # --- Section 2: Het Script (Cheat Sheet) ---
    doc.add_heading('2. Script & Spiekbriefje', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False
    
    # Set column widths
    table.columns[0].width = Inches(1.0) # Slide
    table.columns[1].width = Inches(4.5) # Kernboodschap & Tekst
    table.columns[2].width = Inches(0.8) # Tijd

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Slide'
    hdr_cells[1].text = 'Kernboodschap & Wat te vertellen'
    hdr_cells[2].text = 'Tijd'

    # Helper function to add rows
    def add_row(slide_title, key_message, bullet_points, time):
        row_cells = table.add_row().cells
        
        # Slide Title
        row_cells[0].text = slide_title
        row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Content
        p_msg = row_cells[1].add_paragraph()
        run_msg = p_msg.add_run(f"KERN: {key_message}")
        run_msg.bold = True
        run_msg.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        
        for point in bullet_points:
            row_cells[1].add_paragraph(point, style='List Bullet')
            
        # Time
        row_cells[2].text = time
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    # --- Content Rows ---

    add_row(
        "1. Titel",
        "Welkom & Introductie.",
        [
            "Welkom heten (Hugo, begeleiders, collega's).",
            "Kort voorstellen: Stijn van der Neut, student HBO-ICT.",
            "Titel toelichten: Vandaag neem ik jullie mee in de transformatie van AVE Consultancy."
        ],
        "0:30"
    )

    add_row(
        "2. Agenda",
        "Structuur bieden.",
        [
            "Kort de punten nalopen.",
            "Benadrukken: Eerst de business context, dan de techniek, dan persoonlijke groei.",
            "Meld dat vragen aan het einde mogen (of tussendoor, wat je fijn vindt)."
        ],
        "0:30"
    )

    add_row(
        "3. Situatie",
        "De chaos van Excel/Dropbox.",
        [
            "AVE is een ambitieus bureau, maar de systemen liepen achter.",
            "Beschrijf de oude situatie: 'Bestanden in Dropbox, lijsten in Excel, communicatie via losse mails'.",
            "Het gevolg: Geen inzicht. Wie heeft welke kandidaat gesproken? Alles zat in hoofden van mensen."
        ],
        "2:00"
    )

    add_row(
        "4. Opdracht",
        "SaaS & MVP.",
        [
            "De vraag van Adriaan: 'Bouw een fundering voor de toekomst'.",
            "Niet zomaar een database, maar een SaaS-platform (Software as a Service).",
            "Scope: 20 weken. Focus op de kern: Relaties (CRM) en Kandidaten.",
            "Financiën en App vallen buiten scope."
        ],
        "1:00"
    )

    add_row(
        "5. Probleem",
        "Waarom is dit erg?",
        [
            "Business pijn: Handmatig 3500 CV's verwerken kost maanden.",
            "Risico: GDPR (AVG). Excel sheetjes mailen is niet veilig.",
            "Technisch: Geen relaties. Je weet in Excel niet dat Kandidaat X bij Klant Y op gesprek is geweest."
        ],
        "1:30"
    )

    add_row(
        "6. Onderzoek",
        "Waarom niet kopen?",
        [
            "Belangrijkste slide voor school (Software Adviseren/Analyseren).",
            "Ik heb gekeken naar Bullhorn (te duur/complex) en Recruitee (focus op HR, niet bureaus).",
            "Conclusie Gap-analyse: Er was niets dat én betaalbaar was, én specifiek voor bureaus, én SaaS-ready.",
            "Daarom: Maatwerk (Build vs Buy beslissing)."
        ],
        "2:30"
    )

    add_row(
        "7. Oplossing",
        "De Tech Stack.",
        [
            "High-level overview.",
            "Backend: Laravel (PHP) - Bewezen, veilig, snel.",
            "Frontend: React - Modern, snel, app-gevoel.",
            "Opslag: Cloudflare R2 - Goedkoper dan AWS, sneller dan lokale disk."
        ],
        "1:00"
    )

    add_row(
        "8. Multi-Tenancy",
        "Technische Diepgang 1.",
        [
            "Hoe garanderen we veiligheid als we meerdere klanten op 1 systeem hebben?",
            "Strategie: 'Database per Tenant'.",
            "Leg uit: Klant A heeft Database A. Klant B heeft Database B.",
            "Fysiek gescheiden. Zelfs als de code faalt, kan Klant A nooit data van Klant B zien."
        ],
        "2:30"
    )

    add_row(
        "9. AI Import",
        "Technische Diepgang 2 (Wow-factor).",
        [
            "Probleem: Die 3500 oude CV's.",
            "Oplossing: AI (Gemini 3 Pro) leest de CV's.",
            "Demo-achtig vertellen: 'Het systeem pakt een PDF, leest hem, snapt wat een Skill is, en stopt het in de database'.",
            "Winst: Van 15 min per CV naar secondenwerk."
        ],
        "2:30"
    )

    add_row(
        "10. Tegenslag",
        "Reflectie & Eerlijkheid.",
        [
            "Het moment van de 'Crash': Dataverlies door geen backups.",
            "Eerlijk zijn: Ik schoot in de stress ('Oestergedrag').",
            "De wending: Ik heb het eerlijk opgebiecht en direct een oplossing gebouwd (Automated Backups).",
            "Les: Fouten maken mag, verzwijgen niet."
        ],
        "2:00"
    )

    add_row(
        "11. Prof. Groei",
        "Van Student naar Professional.",
        [
            "Begin: Afwachtend. 'Zeg maar wat ik moet doen'.",
            "Einde: Proactief. 'Ik heb een plan gemaakt voor de migratie'.",
            "Refereer aan feedback Hugo/Adriaan: 'Strategisch partner'."
        ],
        "1:30"
    )

    add_row(
        "12. Toekomst",
        "Hoe nu verder?",
        [
            "Het stopt hier niet.",
            "Nu: Livegang MVP.",
            "Straks: 'Dogfooding' (Zelf gebruiken) en Outlook integratie.",
            "Droom: Dit platform verkopen aan andere bureaus."
        ],
        "1:00"
    )

    add_row(
        "13. Conclusie",
        "Afronding.",
        [
            "Samenvatten: We gingen van chaos naar structuur.",
            "Ik heb laten zien dat ik kan Analyseren, Ontwerpen en Bouwen.",
            "Bedankje richting Adriaan/Hugo voor de kans.",
            "Vragenronde openen."
        ],
        "0:30"
    )

    doc.add_page_break()

    # --- Section 3: Verwachte Vragen (Q&A) ---
    doc.add_heading('3. Verwachte Vragen (Q&A Voorbereiding)', level=1)
    
    qa_list = [
        ("Waarom heb je niet gewoon Salesforce gebruikt?", 
         "Dat heb ik onderzocht. Salesforce is geweldig, maar de licentiekosten voor een starter zijn hoog en de implementatietijd is lang. Voor de specifieke wensen van AVE (snel, simpel, bureau-gericht) was maatwerk op lange termijn goedkoper en effectiever."),
        
        ("Is AI wel veilig met persoonsgegevens?", 
         "Goede vraag. We gebruiken de Enterprise API van Google (Vertex AI/Gemini). De data wordt verwerkt in Europa (regio europe-west4) en Google gebruikt deze data *niet* om hun modellen te trainen. Dit is contractueel vastgelegd."),
        
        ("Wat gebeurt er als je weggaat? Wie onderhoudt dit?", 
         "De code is volledig gedocumenteerd en gebouwd op standaarden (Laravel/React). Elke professionele PHP-ontwikkelaar kan dit overnemen. Daarnaast ligt er een technische overdrachtsdocumentatie."),
        
        ("Waarom Database-per-tenant? Dat is toch duur?", 
         "In opslagruimte valt dat mee, structuur is klein. Het levert vooral enorme veiligheidswinst op. Bij één gedeelde database is één vergeten 'WHERE client_id = ...' al een datalek. Nu is dat fysiek onmogelijk.")
    ]

    for question, answer in qa_list:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"Q: {question}")
        run_q.bold = True
        
        p_a = doc.add_paragraph()
        p_a.add_run(f"A: {answer}")
        
        doc.add_paragraph() # Spacer

    # Save
    filename = "Draaiboek_Eindpresentatie_AVE_CRM.docx"
    doc.save(filename)
    print(f"Successfully generated '{filename}'")

if __name__ == "__main__":
    create_playbook()
